from docx import Document

def apply_text_formatting(src_run, tgt_run):
    tgt_run.font.bold = src_run.font.bold
    tgt_run.font.italic = src_run.font.italic
    tgt_run.font.underline = src_run.font.underline
    tgt_run.font.strike = src_run.font.strike
    tgt_run.font.size = src_run.font.size
    if src_run.font.color.rgb is not None:
        tgt_run.font.color.rgb = src_run.font.color.rgb
    tgt_run.font.name = src_run.font.name

def apply_paragraph_formatting(source_paragraph, target_paragraph):
    target_paragraph.alignment = source_paragraph.alignment
    target_paragraph.paragraph_format.space_before = source_paragraph.paragraph_format.space_before
    target_paragraph.paragraph_format.space_after = source_paragraph.paragraph_format.space_after
    target_paragraph.paragraph_format.line_spacing = source_paragraph.paragraph_format.line_spacing
    target_paragraph.paragraph_format.left_indent = source_paragraph.paragraph_format.left_indent

def copy_header_footer_formatting(original_doc, translated_doc):
    for orig_section, trans_section in zip(original_doc.sections, translated_doc.sections):
        # Copy header formatting
        for orig_header, trans_header in zip(orig_section.header.paragraphs, trans_section.header.paragraphs):
            apply_paragraph_formatting(orig_header, trans_header)
            for orig_run, trans_run in zip(orig_header.runs, trans_header.runs):
                apply_text_formatting(orig_run, trans_run)

        # Copy footer formatting
        for orig_footer, trans_footer in zip(orig_section.footer.paragraphs, trans_section.footer.paragraphs):
            apply_paragraph_formatting(orig_footer, trans_footer)
            for orig_run, trans_run in zip(orig_footer.runs, trans_footer.runs):
                apply_text_formatting(orig_run, trans_run)

def copy_formatting(original_path, translated_path, output_path):
    original_doc = Document(original_path)
    translated_doc = Document(translated_path)

    # Copy formatting for main document content
    for orig_para, trans_para in zip(original_doc.paragraphs, translated_doc.paragraphs):
        apply_paragraph_formatting(orig_para, trans_para)
        for orig_run, trans_run in zip(orig_para.runs, trans_para.runs):
            apply_text_formatting(orig_run, trans_run)

    # Copy formatting for headers and footers
    copy_header_footer_formatting(original_doc, translated_doc)

    translated_doc.save(output_path)


def count_words_in_text(text):
    # Counts the words in a given text segment
    return len(text.split())


def count_words_in_docx(file_path):
    doc = Document(file_path)
    word_count = 0

    # Count words in the main document
    for para in doc.paragraphs:
        word_count += count_words_in_text(para.text)

    # Count words in headers and footers
    for section in doc.sections:
        for header in section.header.paragraphs:
            word_count += count_words_in_text(header.text)
        for footer in section.footer.paragraphs:
            word_count += count_words_in_text(footer.text)

    # Count words in tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    word_count += count_words_in_text(para.text)



    return word_count
